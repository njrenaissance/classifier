"""Text extraction (A2).

Extracts plain text from documents for classification. Each supported format
has its own extractor (the Strategy pattern) that reads from a binary stream;
:func:`extract_text` dispatches on the file suffix through a registry, and
:func:`extract_text_from_bytes` dispatches on the MIME type into the same
registry, so adding a format is a single registration.

The MIME entry point exists for the v2 cloud pipeline (ADR-0015): the processor
downloads a driveItem's bytes into memory and extracts without touching disk,
selecting the strategy from the queue message's ``mime_type`` — no first
download needed to pick it.

Supported now: PDF (via ``pypdf``), DOCX (via ``python-docx``), and plain-text
formats — ``.txt``, ``.json``, ``.yml``/``.yaml``, ``.md``, ``.csv``, ``.xml``,
plus any ``text/*`` MIME type — decoded as raw text (ADR-0021). Legacy binary
``.doc`` is intentionally deferred — ADR-0006 flags it as the weak link with no
reliable pure-Python story — and, like any other unrecognised suffix or MIME
type, is rejected with :class:`~errors.UnsupportedFormatError` rather than
skipped silently.

Extraction never swallows a failure: a missing file, an unreadable stream, or a
corrupt document is surfaced as :class:`~errors.ExtractionError`, chained
(``raise ... from``) from the underlying library error so the root cause
survives in the traceback.
"""

from io import BytesIO
from pathlib import Path
from typing import BinaryIO, Protocol
from zipfile import BadZipFile

from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from pypdf import PdfReader
from pypdf.errors import PyPdfError

from errors import ExtractionError, UnsupportedFormatError


class TextExtractor(Protocol):
    """A strategy that turns one document's bytes into plain text."""

    def extract(self, stream: BinaryIO) -> str: ...


class PdfTextExtractor:
    """Extract text from a PDF page by page using ``pypdf``."""

    def extract(self, stream: BinaryIO) -> str:
        try:
            reader = PdfReader(stream)
            pages = [page.extract_text() or "" for page in reader.pages]
        except (PyPdfError, OSError, ValueError) as err:
            raise ExtractionError("Cannot extract text from PDF document") from err
        return _join_nonblank(pages)


class DocxTextExtractor:
    """Extract text from a DOCX using ``python-docx`` (paragraphs then tables)."""

    def extract(self, stream: BinaryIO) -> str:
        try:
            document = Document(stream)
        except (PackageNotFoundError, BadZipFile, OSError, ValueError) as err:
            raise ExtractionError("Cannot extract text from DOCX document") from err
        lines = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                lines.extend(cell.text for cell in row.cells)
        return _join_nonblank(lines)


class PlainTextExtractor:
    """Decode a text document's bytes to a string (raw, format-agnostic).

    JSON/YAML/XML are not structurally parsed — the classifier wants the textual
    content, and a raw decode keeps one strategy serving every text format
    (ADR-0021). Decodes ``utf-8-sig`` (transparently stripping a byte-order mark
    and handling plain UTF-8); on non-UTF-8 bytes it falls back to ``latin-1``,
    which maps every byte and never raises, so extraction always yields text.
    """

    def extract(self, stream: BinaryIO) -> str:
        try:
            data = stream.read()
        except OSError as err:
            raise ExtractionError("Cannot read plain-text document") from err
        try:
            return data.decode("utf-8-sig")
        except UnicodeDecodeError:
            return data.decode("latin-1")


_PLAIN_TEXT = PlainTextExtractor()  # one shared instance backs every text suffix

_EXTRACTORS: dict[str, TextExtractor] = {
    ".pdf": PdfTextExtractor(),
    ".docx": DocxTextExtractor(),
    ".txt": _PLAIN_TEXT,
    ".json": _PLAIN_TEXT,
    ".yml": _PLAIN_TEXT,
    ".yaml": _PLAIN_TEXT,
    ".md": _PLAIN_TEXT,
    ".csv": _PLAIN_TEXT,
    ".xml": _PLAIN_TEXT,
}

# Canonical MIME per suffix — the single source of truth for the suffix↔MIME
# mapping, and the value :func:`mime_type_for_suffix` stamps on a local file.
_SUFFIX_TO_MIME: dict[str, str] = {
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".txt": "text/plain",
    ".json": "application/json",
    ".yml": "application/yaml",
    ".yaml": "application/yaml",
    ".md": "text/markdown",
    ".csv": "text/csv",
    ".xml": "application/xml",
}

# MIME→suffix for the bytes path: the canonical map inverted, then augmented with
# alias MIME types that share a suffix. ``.yml``/``.yaml`` collapse to one entry —
# harmless, both suffixes resolve to the same plain-text strategy in ``_EXTRACTORS``.
_MIME_TO_SUFFIX: dict[str, str] = {
    **{mime: suffix for suffix, mime in _SUFFIX_TO_MIME.items()},
    "application/x-yaml": ".yaml",
    "text/yaml": ".yaml",
    "text/xml": ".xml",
}

# Any ``text/*`` MIME with no explicit entry above is treated as plain text — a
# catch-all for whatever text subtype SharePoint returns (ADR-0021).
_TEXT_MIME_PREFIX = "text/"


def supported_suffixes() -> frozenset[str]:
    """The lower-cased file suffixes that have a registered extractor."""
    return frozenset(_EXTRACTORS)


def supported_mime_types() -> frozenset[str]:
    """The MIME types that map to a registered extractor."""
    return frozenset(_MIME_TO_SUFFIX)


def mime_type_for_suffix(suffix: str) -> str | None:
    """Return the canonical MIME type for a file ``suffix``, or ``None`` if unsupported.

    Reads the canonical :data:`_SUFFIX_TO_MIME`, so the suffix↔MIME mapping stays
    owned in this one module. The filesystem source (ADR-0020) uses it to stamp a
    driveItem-less local file's ``mime_type`` from its extension, which the
    processor's :func:`extract_text_from_bytes` then dispatches on unchanged.
    """
    return _SUFFIX_TO_MIME.get(suffix.lower())


def extract_text(path: Path) -> str:
    """Extract plain text from ``path``, dispatching on its file suffix.

    Raises :class:`~errors.UnsupportedFormatError` for an unregistered suffix
    (including legacy ``.doc``) and :class:`~errors.ExtractionError` when a
    supported file cannot be read.
    """
    extractor = _EXTRACTORS.get(path.suffix.lower())
    if extractor is None:
        supported = ", ".join(sorted(supported_suffixes()))
        raise UnsupportedFormatError(f"No text extractor for {path.suffix!r} file: {path} (supported: {supported})")
    try:
        with path.open("rb") as stream:
            return extractor.extract(stream)
    except OSError as err:
        raise ExtractionError(f"Cannot read document file: {path}") from err


def extract_text_from_bytes(data: bytes, mime_type: str) -> str:
    """Extract plain text from in-memory ``data``, dispatching on ``mime_type``.

    The processor path (ADR-0015): the downloaded bytes are classified without
    touching disk. Raises :class:`~errors.UnsupportedFormatError` for a MIME type
    with no registered extractor and :class:`~errors.ExtractionError` when the
    bytes cannot be parsed.
    """
    normalized = mime_type.split(";", 1)[0].strip().lower()
    suffix = _MIME_TO_SUFFIX.get(normalized)
    if suffix is None and normalized.startswith(_TEXT_MIME_PREFIX):
        suffix = ".txt"
    extractor = _EXTRACTORS[suffix] if suffix is not None else None
    if extractor is None:
        supported = ", ".join(sorted(supported_mime_types()))
        raise UnsupportedFormatError(f"No text extractor for MIME type {mime_type!r} (supported: {supported})")
    return extractor.extract(BytesIO(data))


def _join_nonblank(parts: list[str]) -> str:
    """Join stripped, non-blank text fragments with single newlines."""
    return "\n".join(stripped for part in parts if (stripped := part.strip()))
