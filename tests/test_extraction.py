from pathlib import Path

import pytest
from docx import Document

from errors import AppError, ExtractionError, UnsupportedFormatError
from extraction import (
    extract_text,
    extract_text_from_bytes,
    mime_type_for_suffix,
    supported_mime_types,
    supported_suffixes,
)

pytestmark = pytest.mark.unit

_PDF_MIME = "application/pdf"
_DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
_TXT_MIME = "text/plain"
_JSON_MIME = "application/json"
_YAML_MIME = "application/yaml"
_MD_MIME = "text/markdown"
_CSV_MIME = "text/csv"
_XML_MIME = "application/xml"

_TEXT_SUFFIXES = (".txt", ".json", ".yml", ".yaml", ".md", ".csv", ".xml")


def _build_pdf(text: str) -> bytes:
    """Build a minimal single-page PDF whose only content is ``text``.

    Hand-rolled (rather than pulling in a PDF-writing dependency) so the real
    ``pypdf`` extraction path is exercised end-to-end against a valid file.
    """
    escaped = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    content = f"BT /F1 24 Tf 72 720 Td ({escaped}) Tj ET".encode("latin-1")
    bodies = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Resources << /Font << /F1 5 0 R >> >> /Contents 4 0 R >>",
        b"<< /Length " + str(len(content)).encode() + b" >>\nstream\n" + content + b"\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    out = bytearray(b"%PDF-1.4\n")
    offsets: list[int] = []
    for number, body in enumerate(bodies, start=1):
        offsets.append(len(out))
        out += f"{number} 0 obj\n".encode() + body + b"\nendobj\n"
    xref_pos = len(out)
    out += f"xref\n0 {len(bodies) + 1}\n".encode()
    out += b"0000000000 65535 f \n"
    out += b"".join(f"{offset:010d} 00000 n \n".encode() for offset in offsets)
    out += f"trailer\n<< /Size {len(bodies) + 1} /Root 1 0 R >>\nstartxref\n{xref_pos}\n%%EOF".encode()
    return bytes(out)


def _write_docx(path: Path, paragraphs: tuple[str, ...], table_row: tuple[str, ...] | None = None) -> None:
    document = Document()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    if table_row is not None:
        table = document.add_table(rows=1, cols=len(table_row))
        for cell, value in zip(table.rows[0].cells, table_row, strict=True):
            cell.text = value
    document.save(str(path))


def test_extracts_text_from_pdf(tmp_path: Path):
    pdf = tmp_path / "invoice.pdf"
    pdf.write_bytes(_build_pdf("INVOICE 4521 total due"))
    assert "INVOICE 4521 total due" in extract_text(pdf)


def test_extracts_paragraphs_and_tables_from_docx(tmp_path: Path):
    docx = tmp_path / "contract.docx"
    _write_docx(docx, ("Master Services Agreement", "Between Acme and Globex"), table_row=("Term", "12 months"))
    result = extract_text(docx)
    assert "Master Services Agreement" in result
    assert "Between Acme and Globex" in result
    assert "12 months" in result


def test_docx_extraction_drops_blank_paragraphs(tmp_path: Path):
    docx = tmp_path / "spaced.docx"
    _write_docx(docx, ("First", "", "   ", "Second"))
    assert extract_text(docx) == "First\nSecond"


@pytest.mark.parametrize("suffix", _TEXT_SUFFIXES)
def test_plain_text_is_decoded_verbatim(tmp_path: Path, suffix: str):
    content = '{\n  "note": "line one"\n}\n\n  trailing indent  '
    path = tmp_path / f"doc{suffix}"
    path.write_bytes(content.encode("utf-8"))  # bytes, so no newline translation
    # Raw decode: whitespace, blank lines and indentation are preserved as-is.
    assert extract_text(path) == content


def test_plain_text_strips_utf8_bom(tmp_path: Path):
    path = tmp_path / "bom.txt"
    path.write_bytes("hello world".encode("utf-8-sig"))
    assert extract_text(path) == "hello world"


def test_plain_text_falls_back_to_latin1_on_non_utf8_bytes(tmp_path: Path):
    path = tmp_path / "resume.txt"
    path.write_bytes(b"caf\xe9")  # 0xE9 is 'é' in latin-1, invalid as UTF-8
    assert extract_text(path) == "café"


def test_empty_plain_text_file_yields_empty_string(tmp_path: Path):
    path = tmp_path / "empty.txt"
    path.write_bytes(b"")
    assert extract_text(path) == ""


@pytest.mark.parametrize("suffix", [".pdf", ".PDF", ".docx", ".DoCx"])
def test_dispatch_is_case_insensitive_on_suffix(tmp_path: Path, suffix: str):
    path = tmp_path / f"doc{suffix}"
    if suffix.lower() == ".pdf":
        path.write_bytes(_build_pdf("hello"))
    else:
        _write_docx(path, ("hello",))
    assert "hello" in extract_text(path)


@pytest.mark.parametrize(
    "name",
    [
        pytest.param("legacy.doc", id="legacy_doc_is_deferred"),
        pytest.param("archive.zip", id="unrelated_binary"),
        pytest.param("noextension", id="no_suffix"),
    ],
)
def test_unsupported_suffix_is_rejected_explicitly(tmp_path: Path, name: str):
    path = tmp_path / name
    path.write_bytes(b"whatever")
    with pytest.raises(UnsupportedFormatError):
        extract_text(path)


@pytest.mark.parametrize("name", ["missing.pdf", "missing.docx"])
def test_missing_file_is_surfaced_as_extraction_error(tmp_path: Path, name: str):
    with pytest.raises(ExtractionError):
        extract_text(tmp_path / name)


@pytest.mark.parametrize("name", ["broken.pdf", "broken.docx"])
def test_corrupt_file_is_surfaced_as_extraction_error(tmp_path: Path, name: str):
    path = tmp_path / name
    path.write_bytes(b"this is not a valid document")
    with pytest.raises(ExtractionError):
        extract_text(path)


def test_corrupt_file_error_chains_original_cause(tmp_path: Path):
    path = tmp_path / "broken.pdf"
    path.write_bytes(b"this is not a valid document")
    with pytest.raises(ExtractionError) as excinfo:
        extract_text(path)
    assert excinfo.value.__cause__ is not None


def test_unsupported_format_error_is_an_app_error(tmp_path: Path):
    with pytest.raises(AppError):
        extract_text(tmp_path / "legacy.doc")


def test_supported_suffixes_include_documents_and_plain_text():
    assert supported_suffixes() == frozenset({".pdf", ".docx", *_TEXT_SUFFIXES})


def test_supported_mime_types_include_documents_and_plain_text():
    assert supported_mime_types() == frozenset(
        {
            _PDF_MIME,
            _DOCX_MIME,
            _TXT_MIME,
            _JSON_MIME,
            _YAML_MIME,
            _MD_MIME,
            _CSV_MIME,
            _XML_MIME,
            "application/x-yaml",
            "text/yaml",
            "text/xml",
        }
    )


@pytest.mark.parametrize(
    ("suffix", "expected"),
    [
        pytest.param(".pdf", _PDF_MIME, id="pdf"),
        pytest.param(".docx", _DOCX_MIME, id="docx"),
        pytest.param(".PDF", _PDF_MIME, id="uppercase_normalised"),
        pytest.param(".doc", None, id="legacy_doc_unsupported"),
        pytest.param(".txt", _TXT_MIME, id="plain_text"),
        pytest.param(".json", _JSON_MIME, id="json"),
        pytest.param(".yml", _YAML_MIME, id="yml"),
        pytest.param(".yaml", _YAML_MIME, id="yaml"),
        pytest.param(".md", _MD_MIME, id="markdown"),
        pytest.param(".csv", _CSV_MIME, id="csv"),
        pytest.param(".xml", _XML_MIME, id="xml"),
        pytest.param("", None, id="empty_suffix"),
    ],
)
def test_mime_type_for_suffix(suffix: str, expected: str | None):
    assert mime_type_for_suffix(suffix) == expected


def _write_pdf(path: Path) -> None:
    path.write_bytes(_build_pdf("INVOICE 4521 total due"))


def _write_docx_sample(path: Path) -> None:
    _write_docx(path, ("Master Services Agreement", "Between Acme and Globex"), table_row=("Term", "12 months"))


def _write_json_sample(path: Path) -> None:
    path.write_text('{"customer": "Acme", "total": 4521}\n', encoding="utf-8")


@pytest.mark.parametrize(
    ("write_file", "name", "mime_type"),
    [
        pytest.param(_write_pdf, "invoice.pdf", _PDF_MIME, id="pdf"),
        pytest.param(_write_docx_sample, "contract.docx", _DOCX_MIME, id="docx"),
        pytest.param(_write_json_sample, "invoice.json", _JSON_MIME, id="json"),
    ],
)
def test_bytes_extraction_matches_path_extraction(tmp_path: Path, write_file, name: str, mime_type: str):
    path = tmp_path / name
    write_file(path)
    data = path.read_bytes()

    from_bytes = extract_text_from_bytes(data, mime_type)

    assert from_bytes != ""
    assert from_bytes == extract_text(path)


@pytest.mark.parametrize(
    "mime_type",
    [
        pytest.param(_TXT_MIME, id="text_plain"),
        pytest.param(_MD_MIME, id="text_markdown_explicit"),
        pytest.param("text/x-anything", id="unlisted_text_catch_all"),
        pytest.param("text/csv; charset=utf-8", id="text_with_charset_parameter"),
    ],
)
def test_any_text_mime_type_extracts_as_plain_text(mime_type: str):
    assert extract_text_from_bytes(b"hello world", mime_type) == "hello world"


@pytest.mark.parametrize(
    "mime_type",
    [
        pytest.param("APPLICATION/PDF", id="uppercase"),
        pytest.param("application/pdf; charset=binary", id="with_parameter"),
    ],
)
def test_bytes_dispatch_normalizes_mime_type(mime_type: str):
    assert "hello" in extract_text_from_bytes(_build_pdf("hello"), mime_type)


@pytest.mark.parametrize(
    "mime_type",
    [
        pytest.param("application/msword", id="legacy_doc"),
        pytest.param("application/zip", id="unrelated_binary"),
        pytest.param("", id="empty"),
    ],
)
def test_unsupported_mime_type_is_rejected_explicitly(mime_type: str):
    with pytest.raises(UnsupportedFormatError):
        extract_text_from_bytes(b"whatever", mime_type)


def test_unsupported_mime_type_error_is_an_app_error():
    with pytest.raises(AppError):
        extract_text_from_bytes(b"whatever", "application/msword")


@pytest.mark.parametrize("mime_type", [_PDF_MIME, _DOCX_MIME])
def test_corrupt_byte_stream_is_surfaced_as_extraction_error(mime_type: str):
    with pytest.raises(ExtractionError):
        extract_text_from_bytes(b"this is not a valid document", mime_type)


def test_corrupt_byte_stream_error_chains_original_cause():
    with pytest.raises(ExtractionError) as excinfo:
        extract_text_from_bytes(b"this is not a valid document", _PDF_MIME)
    assert excinfo.value.__cause__ is not None
